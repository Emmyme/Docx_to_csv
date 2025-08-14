import os
import csv
import win32com.client
import pythoncom
from docx import Document

# Keywords to search for in the documents
FIELD_KEYWORDS = {
    "KEYWORD_1": ["EXAMPLE_1", "EXAMPLE_2", "EXAMPLE_3", "EXAMPLE_4", "EXAMPLE_5", "EXAMPLE_6"],
    "KEYWORD_2": ["EXAMPLE_1", "EXAMPLE_2", "EXAMPLE_3", "EXAMPLE_4", "EXAMPLE_5", "EXAMPLE_6"],
    "KEYWORD_3": ["EXAMPLE_1", "EXAMPLE_2", "EXAMPLE_3", "EXAMPLE_4", "EXAMPLE_5", "EXAMPLE_6"],
}

# Only for .doc files
def convert_to_docx(doc_path):
    # Initialize COM (needed for Word automation)
    pythoncom.CoInitialize()
    
    # Get absolute paths (needed for pywin32)
    abs_path = os.path.abspath(doc_path)
    docx_path = os.path.splitext(abs_path)[0] + '.docx'
    
    if os.path.exists(docx_path):
        return docx_path
    
    word = None
    doc = None
    
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        doc = word.Documents.Open(abs_path)
        
        # Save as docx
        doc.SaveAs2(docx_path, FileFormat=16)

        doc.Close(SaveChanges=False)
        doc = None
        
        # Delete the original .doc file
        if os.path.exists(abs_path):
            try:
                os.remove(abs_path)
            except:
                pass

        return docx_path
    
    except:
        return None
    finally:

        if doc:
            try:
                doc.Close(SaveChanges=False)
            except:
                pass
        if word:
            try:
                word.Quit()
            except:
                pass
        pythoncom.CoUninitialize()

def find_field(text_content):
    extracted = {field: "" for field in FIELD_KEYWORDS}

    paragraphs = text_content.split('\n')
    
    for text in paragraphs:
        if not text.strip(): 
            continue
            
        info = text.strip().lower()
        
        for field, keywords in FIELD_KEYWORDS.items():
            for keyword in keywords:
                if keyword.lower() in info:
                    parts = text.split(":")
                    if len(parts) > 1:
                        extracted[field] = parts[1].strip()
                    else:
                        extracted[field] = text.strip()
                    break 

    return extracted

def extract_from_doc(doc_path):
    # Convert to docx if needed
    if doc_path.endswith('.doc'):
        docx_path = convert_to_docx(doc_path)
        if not docx_path:
            return {field: "" for field in FIELD_KEYWORDS}
        doc_path = docx_path
    
    try:
        # Read the document
        doc = Document(doc_path)
        
        # Paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        text_content = "\n".join(paragraphs)
        
        # Process the text 
        return find_field(text_content)
    except:
        return {field: "" for field in FIELD_KEYWORDS}

def extract_from_filename(filename):
    name = os.path.splitext(filename)[0]
    parts = name.split(" ")

    if len(parts) >= 3:
        return {
            "NAME": f"{parts[0]} {parts[1]}",
            "DESCRIPTION": " ".join(parts[2:]),
        }
    else:
        return {
            "NAME": "",
            "DESCRIPTION": "",
        }

def main(folder_path, output_csv):
    all_data = []
    
    for filename in os.listdir(folder_path):
        if filename.endswith(".doc"):
            full_path = os.path.join(folder_path, filename)
            convert_to_docx(full_path)

    # Process the files
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            full_path = os.path.join(folder_path, filename)
            
            # Get data from document
            data = extract_from_doc(full_path)
            # Get data from filename
            file_info = extract_from_filename(filename)
            # Combine them
            data.update(file_info)

            if 'filename' in data:
                del data['filename']
            
            all_data.append(data)

    fieldnames = ["NAME", "DESCRIPTION"] + list(FIELD_KEYWORDS.keys())

    # Write to CSV
    with open(output_csv, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter=';')
        writer.writeheader()
        writer.writerows(all_data)


if __name__ == "__main__":
    folder_path = "./docs"
    output_csv = "output.csv"
    main(folder_path, output_csv)